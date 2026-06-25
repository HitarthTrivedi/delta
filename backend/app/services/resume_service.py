"""
Resume Service — delta Career OS
Handles: generation, suggestion diffing, ATS optimization, and .docx export.
"""
from __future__ import annotations

import datetime
import io
import json
import re
import uuid
from typing import Any

from sqlalchemy.orm import Session

from app.models import (
    CareerMemoryProfile,
    JourneyEvent,
    MarketSnapshot,
    ResumeProfile,
    SkillNode,
    User,
)

# ─── helpers ──────────────────────────────────────────────────────────────────

def _j(value: Any, fallback: Any = None) -> Any:
    if value is None:
        return fallback
    if isinstance(value, (dict, list)):
        return value
    try:
        return json.loads(value)
    except Exception:
        return fallback


def _dump(value: Any) -> str:
    return json.dumps(value, ensure_ascii=True)


# ─── Generation ───────────────────────────────────────────────────────────────

def build_resume_from_profile(user_id: str, db: Session) -> dict:
    """
    Construct a structured resume JSON from the user's Career OS profile.
    Returns the structured_data dict.
    """
    user: User | None = db.query(User).filter(User.id == user_id).first()
    if not user:
        raise ValueError("User not found")

    memory: CareerMemoryProfile | None = (
        db.query(CareerMemoryProfile).filter(CareerMemoryProfile.user_id == user_id).first()
    )
    skills_db = db.query(SkillNode).filter(SkillNode.user_id == user_id).all()

    identity = _j(memory.identity if memory else None, {})
    ambitions = _j(memory.ambitions if memory else None, {})
    capabilities = _j(memory.capabilities if memory else None, {})
    evidence = _j(memory.evidence if memory else None, {})

    name = identity.get("name") or user.name or "Candidate"
    email = identity.get("email") or user.email or ""
    target_role = ambitions.get("target_role") or user.target_role or "Software Engineer"
    education_stage = identity.get("education_stage") or user.current_role or ""
    background_summary = identity.get("background_summary") or ""

    # ── Skills ──
    skill_names: list[str] = []
    seen: set[str] = set()
    # High-confidence skills first (proficiency ≥ 6 or evidence verified)
    for skill in sorted(skills_db, key=lambda s: -(s.proficiency or 0)):
        key = skill.name.strip().lower()
        if key and key not in seen:
            seen.add(key)
            skill_names.append(skill.name.strip())

    # Fill from capabilities if skill list is sparse
    cap_skills = capabilities.get("current_skills") or []
    if isinstance(cap_skills, list):
        for s in cap_skills:
            key = str(s).strip().lower()
            if key and key not in seen:
                seen.add(key)
                skill_names.append(str(s).strip())

    # ── Experience / Projects from evidence map ──
    evidence_projects = evidence.get("projects") or []
    experience_rows: list[dict] = []
    project_rows: list[dict] = []
    for proj in evidence_projects:
        if isinstance(proj, dict):
            entry = {
                "title": proj.get("title") or proj.get("name") or "Project",
                "description": proj.get("description") or proj.get("summary") or "",
                "tech": proj.get("tech") or proj.get("technologies") or [],
                "url": proj.get("url") or proj.get("github") or "",
                "date": proj.get("date") or "",
            }
            project_rows.append(entry)

    # ── Journey-event derived achievements ──
    completed_events = (
        db.query(JourneyEvent)
        .filter(
            JourneyEvent.user_id == user_id,
            JourneyEvent.event_type.in_(
                ["task_completed", "weekly_task_completed", "milestone_achieved"]
            ),
        )
        .order_by(JourneyEvent.created_at.desc())
        .limit(20)
        .all()
    )
    achievements: list[str] = []
    for evt in completed_events:
        summary = evt.summary or ""
        if summary and len(summary) > 10:
            # Clean up internal prefix
            clean = re.sub(r"^(Completed action item:|Reopened action item:)\s*", "", summary, flags=re.IGNORECASE).strip().strip('"')
            if clean and clean not in achievements:
                achievements.append(clean)

    # ── Summary statement ──
    long_goals = (ambitions.get("long_term_goals") or [])
    goal_text = long_goals[0] if long_goals else target_role
    summary = (
        background_summary
        or f"Results-oriented {target_role} with hands-on experience in "
        f"{', '.join(skill_names[:3]) if skill_names else 'software development'}. "
        f"Motivated to {goal_text.lower() if goal_text else 'deliver impactful solutions'}."
    )

    structured = {
        "contact": {
            "name": name,
            "email": email,
            "role": target_role,
        },
        "summary": summary,
        "skills": skill_names[:20],
        "experience": experience_rows,
        "projects": project_rows,
        "achievements": achievements[:8],
        "education": [
            {
                "degree": education_stage,
                "institution": identity.get("university") or identity.get("college") or "",
                "year": identity.get("graduation_year") or "",
            }
        ],
        "certifications": evidence.get("certificates") or [],
        "_meta": {
            "generated_at": datetime.datetime.utcnow().isoformat(),
            "source": "delta_profile",
        },
    }
    return structured


# ─── ATS Score ────────────────────────────────────────────────────────────────

def _compute_ats_score(structured: dict, market_skills: list[str]) -> float:
    """Ratio of market skills that appear in the resume text."""
    if not market_skills:
        return 0.5
    resume_text = json.dumps(structured).lower()
    hits = sum(1 for s in market_skills if s.lower() in resume_text)
    return round(hits / len(market_skills), 3)


def optimize_for_ats(structured: dict, market_skills: list[str]) -> dict:
    """
    Rewrites bullet points (achievements, project descriptions) to
    incorporate market keywords naturally, and updates the ATS score.
    """
    from app.services.ai_service import generate_json

    if not market_skills:
        return structured

    top_keywords = market_skills[:15]

    # Build a compact resume text for the AI
    resume_text = json.dumps({
        "summary": structured.get("summary", ""),
        "skills": structured.get("skills", []),
        "achievements": structured.get("achievements", []),
        "projects": [
            {"title": p.get("title"), "description": p.get("description")}
            for p in (structured.get("projects") or [])
        ],
    }, indent=2)

    prompt = f"""You are an expert ATS resume optimizer.

Target keywords (from current job market): {json.dumps(top_keywords)}

Current resume content:
{resume_text}

Rewrite the content so that:
1. The summary naturally incorporates 3-5 of the target keywords.
2. Each achievement bullet starts with a strong action verb and mentions at least one keyword if relevant.
3. Project descriptions mention relevant keywords without keyword stuffing.
4. Skills list is reordered to put market-demanded skills first.
5. Language stays truthful — do not invent experience.

Return ONLY a JSON object with keys: summary, achievements (array), projects (array of {{title, description}}), skills (array).
"""
    result = generate_json(prompt, temperature=0.3)

    if isinstance(result, dict):
        if result.get("summary"):
            structured = {**structured, "summary": result["summary"]}
        if isinstance(result.get("achievements"), list):
            structured = {**structured, "achievements": result["achievements"]}
        if isinstance(result.get("skills"), list):
            structured = {**structured, "skills": result["skills"]}
        if isinstance(result.get("projects"), list):
            existing_projects = structured.get("projects") or []
            merged_projects = []
            for orig in existing_projects:
                match = next(
                    (p for p in result["projects"] if p.get("title") == orig.get("title")),
                    None,
                )
                if match:
                    merged_projects.append({**orig, "description": match.get("description", orig.get("description"))})
                else:
                    merged_projects.append(orig)
            structured = {**structured, "projects": merged_projects}

    structured["_ats_optimized"] = True
    structured["_ats_keywords"] = top_keywords
    return structured


# ─── Bi-weekly Suggestions ────────────────────────────────────────────────────

def generate_suggestions(user_id: str, db: Session, resume: ResumeProfile) -> dict:
    """
    Diff the stored resume against new journey events and skill changes
    since last_suggested_at. Returns { to_add: [...], to_remove: [...] }.
    """
    structured = _j(resume.structured_data, {})
    since = resume.last_suggested_at or (resume.updated_at - datetime.timedelta(days=14))

    # ── New completed tasks / achievements ──
    new_events = (
        db.query(JourneyEvent)
        .filter(
            JourneyEvent.user_id == user_id,
            JourneyEvent.event_type.in_(
                ["task_completed", "weekly_task_completed", "milestone_achieved"]
            ),
            JourneyEvent.created_at >= since,
        )
        .order_by(JourneyEvent.created_at.desc())
        .all()
    )

    existing_achievements = set(structured.get("achievements") or [])
    to_add: list[dict] = []

    for evt in new_events:
        summary = evt.summary or ""
        clean = re.sub(
            r"^(Completed action item:|Reopened action item:)\s*", "", summary, flags=re.IGNORECASE
        ).strip().strip('"')
        if clean and len(clean) > 10 and clean not in existing_achievements:
            to_add.append({
                "type": "achievement",
                "value": clean,
                "source": "completed_task",
                "date": str(evt.event_date or evt.created_at.date()),
            })

    # ── New skills ──
    existing_skills = {s.lower() for s in (structured.get("skills") or [])}
    new_skills = (
        db.query(SkillNode)
        .filter(
            SkillNode.user_id == user_id,
            SkillNode.created_at >= since,
            SkillNode.proficiency >= 5,
        )
        .all()
    )
    for skill in new_skills:
        if skill.name.lower() not in existing_skills:
            to_add.append({
                "type": "skill",
                "value": skill.name,
                "source": "skill_node",
                "date": str(skill.created_at.date()),
            })

    # ── Identify items worth removing ──
    # (skills with low market signal vs. market demanded)
    market = (
        db.query(MarketSnapshot)
        .filter(MarketSnapshot.user_id == user_id)
        .order_by(MarketSnapshot.snapshot_date.desc())
        .first()
    )
    demanded_skills: list[str] = []
    if market:
        demanded_skills = _j(market.top_demanded_skills, [])

    to_remove: list[dict] = []
    # Flag skills in resume that don't appear in demanded list and have low uniqueness
    resume_skills = structured.get("skills") or []
    demanded_lower = {s.lower() for s in demanded_skills}
    generic_skills = {"python", "javascript", "git", "html", "css"}  # Keep these always
    for skill in resume_skills:
        sl = skill.lower()
        if sl not in demanded_lower and sl not in generic_skills:
            to_remove.append({
                "type": "skill",
                "value": skill,
                "reason": "Not currently in market demand signals — consider replacing with a more targeted skill.",
            })

    # Flag generic/vague achievement bullets
    for ach in (structured.get("achievements") or []):
        vague_patterns = [
            r"^learned\s",
            r"^studied\s",
            r"^read\s",
            r"^watched\s",
            r"^completed.*course",
        ]
        if any(re.search(pat, ach.lower()) for pat in vague_patterns):
            to_remove.append({
                "type": "achievement",
                "value": ach,
                "reason": "Passive learning statement — ATS prefers action-verb bullet points.",
            })

    return {
        "to_add": to_add[:12],
        "to_remove": to_remove[:8],
        "since": since.isoformat(),
        "generated_at": datetime.datetime.utcnow().isoformat(),
    }


def apply_suggestions(
    structured: dict,
    accepted_adds: list[dict],
    accepted_removes: list[str],
) -> dict:
    """
    Merge accepted additions and deletions into the structured resume.
    accepted_removes: list of values (text strings) to remove.
    """
    result = dict(structured)

    # Apply adds
    new_skills = [item["value"] for item in accepted_adds if item.get("type") == "skill"]
    new_achievements = [item["value"] for item in accepted_adds if item.get("type") == "achievement"]

    existing_skills = list(result.get("skills") or [])
    existing_skill_lower = {s.lower() for s in existing_skills}
    for s in new_skills:
        if s.lower() not in existing_skill_lower:
            existing_skills.append(s)
            existing_skill_lower.add(s.lower())
    result["skills"] = existing_skills[:20]

    existing_achievements = list(result.get("achievements") or [])
    existing_ach_set = {a.lower() for a in existing_achievements}
    for a in new_achievements:
        if a.lower() not in existing_ach_set:
            existing_achievements.append(a)
            existing_ach_set.add(a.lower())
    result["achievements"] = existing_achievements[:10]

    # Apply removes
    removes_lower = {v.lower() for v in accepted_removes}
    result["skills"] = [s for s in result.get("skills", []) if s.lower() not in removes_lower]
    result["achievements"] = [a for a in result.get("achievements", []) if a.lower() not in removes_lower]

    result["_meta"] = {
        **(result.get("_meta") or {}),
        "last_suggestions_applied": datetime.datetime.utcnow().isoformat(),
    }
    return result


# ─── .docx Export ─────────────────────────────────────────────────────────────

def export_to_docx(structured: dict) -> bytes:
    """
    Generate an ATS-friendly .docx file from the structured resume dict.
    Returns the file as bytes.
    """
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    doc = Document()

    # ── Page margins (narrow for ATS) ──
    for section in doc.sections:
        section.top_margin = Inches(0.75)
        section.bottom_margin = Inches(0.75)
        section.left_margin = Inches(0.85)
        section.right_margin = Inches(0.85)

    # ── Default font (ATS-safe) ──
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(10.5)

    contact = structured.get("contact") or {}

    # ── Header: Name ──
    name_para = doc.add_paragraph()
    name_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    name_run = name_para.add_run(contact.get("name") or "Candidate")
    name_run.bold = True
    name_run.font.size = Pt(18)

    # ── Header: Role + Email ──
    role_para = doc.add_paragraph()
    role_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    role_text = " | ".join(filter(None, [contact.get("role"), contact.get("email")]))
    role_run = role_para.add_run(role_text)
    role_run.font.size = Pt(10)

    def _section_heading(title: str):
        para = doc.add_paragraph()
        para.paragraph_format.space_before = Pt(8)
        run = para.add_run(title.upper())
        run.bold = True
        run.font.size = Pt(10.5)
        run.font.color.rgb = RGBColor(0x1A, 0x56, 0xAB)  # professional blue
        # Add a thin horizontal rule via border
        pPr = para._p.get_or_add_pPr()
        pBdr = OxmlElement("w:pBdr")
        bottom = OxmlElement("w:bottom")
        bottom.set(qn("w:val"), "single")
        bottom.set(qn("w:sz"), "6")
        bottom.set(qn("w:space"), "1")
        bottom.set(qn("w:color"), "1A56AB")
        pBdr.append(bottom)
        pPr.append(pBdr)

    # ── Professional Summary ──
    summary = structured.get("summary") or ""
    if summary:
        _section_heading("Professional Summary")
        doc.add_paragraph(summary)

    # ── Skills ──
    skills = structured.get("skills") or []
    if skills:
        _section_heading("Technical Skills")
        # Group into rows of 4 for readability
        chunks = [skills[i:i+4] for i in range(0, len(skills), 4)]
        for chunk in chunks:
            doc.add_paragraph(" • ".join(chunk))

    # ── Experience ──
    experience = structured.get("experience") or []
    if experience:
        _section_heading("Experience")
        for exp in experience:
            p = doc.add_paragraph()
            title_run = p.add_run(f"{exp.get('title', '')} — {exp.get('company', '')}")
            title_run.bold = True
            if exp.get("date"):
                p.add_run(f"  |  {exp['date']}")
            desc = exp.get("description") or ""
            if desc:
                doc.add_paragraph(desc, style="List Bullet")

    # ── Projects ──
    projects = structured.get("projects") or []
    if projects:
        _section_heading("Projects")
        for proj in projects:
            p = doc.add_paragraph()
            title_run = p.add_run(proj.get("title") or "Project")
            title_run.bold = True
            tech = proj.get("tech") or []
            if isinstance(tech, list) and tech:
                p.add_run(f"  |  {', '.join(tech)}")
            elif isinstance(tech, str) and tech:
                p.add_run(f"  |  {tech}")
            desc = proj.get("description") or ""
            if desc:
                doc.add_paragraph(desc, style="List Bullet")

    # ── Achievements / Accomplishments ──
    achievements = structured.get("achievements") or []
    if achievements:
        _section_heading("Key Achievements")
        for ach in achievements:
            doc.add_paragraph(ach, style="List Bullet")

    # ── Education ──
    education = structured.get("education") or []
    if education:
        _section_heading("Education")
        for edu in education:
            degree = edu.get("degree") or ""
            institution = edu.get("institution") or ""
            year = edu.get("year") or ""
            parts = filter(None, [degree, institution, year])
            doc.add_paragraph(" | ".join(parts))

    # ── Certifications ──
    certifications = structured.get("certifications") or []
    if certifications:
        _section_heading("Certifications")
        for cert in certifications:
            if isinstance(cert, dict):
                doc.add_paragraph(cert.get("name") or str(cert), style="List Bullet")
            else:
                doc.add_paragraph(str(cert), style="List Bullet")

    # ── Write to bytes ──
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()


# ─── Persistence helpers ───────────────────────────────────────────────────────

def get_or_create_resume_profile(user_id: str, db: Session) -> ResumeProfile | None:
    return db.query(ResumeProfile).filter(ResumeProfile.user_id == user_id).first()


def save_resume(
    user_id: str,
    db: Session,
    structured: dict,
    raw_text: str = "",
    source: str = "generated",
    market_skills: list[str] | None = None,
) -> ResumeProfile:
    ats_score = _compute_ats_score(structured, market_skills or [])
    existing = db.query(ResumeProfile).filter(ResumeProfile.user_id == user_id).first()
    if existing:
        existing.structured_data = _dump(structured)
        existing.raw_text = raw_text or existing.raw_text
        existing.source = source
        existing.ats_score = ats_score
        existing.updated_at = datetime.datetime.utcnow()
        db.commit()
        db.refresh(existing)
        return existing
    else:
        resume = ResumeProfile(
            id=str(uuid.uuid4()),
            user_id=user_id,
            structured_data=_dump(structured),
            raw_text=raw_text,
            source=source,
            ats_score=ats_score,
        )
        db.add(resume)
        db.commit()
        db.refresh(resume)
        return resume


def get_market_skills(user_id: str, db: Session) -> list[str]:
    market = (
        db.query(MarketSnapshot)
        .filter(MarketSnapshot.user_id == user_id)
        .order_by(MarketSnapshot.snapshot_date.desc())
        .first()
    )
    if not market:
        return []
    return _j(market.top_demanded_skills, [])


def serialize_resume(resume: ResumeProfile) -> dict:
    return {
        "id": resume.id,
        "user_id": resume.user_id,
        "structured_data": _j(resume.structured_data, {}),
        "ats_score": resume.ats_score,
        "source": resume.source,
        "last_suggested_at": resume.last_suggested_at.isoformat() if resume.last_suggested_at else None,
        "updated_at": resume.updated_at.isoformat() if resume.updated_at else None,
    }


def suggestions_are_due(resume: ResumeProfile) -> bool:
    """Returns True if it has been ≥14 days since the last suggestion check."""
    if resume.last_suggested_at is None:
        # Fresh resume — suggest after 14 days from creation
        if resume.updated_at:
            return (datetime.datetime.utcnow() - resume.updated_at).days >= 14
        return False
    return (datetime.datetime.utcnow() - resume.last_suggested_at).days >= 14
